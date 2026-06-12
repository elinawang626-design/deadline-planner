"""Attachment text extraction and deterministic estimate-context selection.

Extraction supports TXT/Markdown (encoding fallback), PDF (pypdf) and DOCX
(python-docx). pypdf / python-docx are imported lazily so the server still
runs without them; affected files are marked ``failed`` instead.

Selection compresses extracted text into the highest-signal fragments under a
character budget. It is a deterministic keyword/position scorer, not document
understanding: fragments keep their source so the external AI can cite them.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

# Character budgets for prompt context (per file and across all files).
PER_FILE_CHAR_BUDGET = 3000
TOTAL_CHAR_BUDGET = 12000

_TEXT_SUFFIXES = {".txt", ".md", ".markdown"}
_ENCODINGS = ("utf-8", "gb18030", "latin-1")

# Signals that a paragraph talks about scope, deliverables or constraints.
_KEYWORDS = (
    "需求", "要求", "交付", "提交", "验收", "评分", "标准", "截止", "依赖",
    "约束", "风险", "必须", "需要", "完成", "实现", "测试", "文档", "数量",
    "requirement", "deliverable", "due", "deadline", "submit", "rubric",
    "must", "should", "acceptance", "constraint", "risk", "depend", "grade",
)
_HEADING_RE = re.compile(r"^(#{1,6}\s+|\d+[\.、]\s*|[一二三四五六七八九十]+[、.]\s*)")
_NUMBER_RE = re.compile(r"\d")


def extract_text(path: Path, display_name: str) -> tuple[str, str | None]:
    """Return (extractionStatus, text). Never raises on unreadable content."""
    suffix = Path(display_name).suffix.lower()
    if suffix in _TEXT_SUFFIXES:
        return _extract_plain(path)
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    return "unsupported", None


def _extract_plain(path: Path) -> tuple[str, str | None]:
    try:
        raw = path.read_bytes()
    except OSError:
        return "failed", None
    for encoding in _ENCODINGS:
        try:
            return "ok", raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return "failed", None


def _extract_pdf(path: Path) -> tuple[str, str | None]:
    try:
        from pypdf import PdfReader
    except ImportError:
        return "failed", None
    try:
        reader = PdfReader(str(path))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception:
        return "failed", None
    text = "\n\n".join(p.strip() for p in pages if p.strip())
    # Scanned PDFs extract to (almost) nothing; keep the attachment but flag it.
    return ("ok", text) if text.strip() else ("failed", None)


def _extract_docx(path: Path) -> tuple[str, str | None]:
    try:
        import docx
    except ImportError:
        return "failed", None
    try:
        document = docx.Document(str(path))
    except Exception:
        return "failed", None
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n\n".join(parts)
    return ("ok", text) if text.strip() else ("failed", None)


# ---- deterministic selection ----


@dataclass(frozen=True)
class Excerpt:
    source: str  # attachment display name
    heading: str  # nearest heading line, "" when none
    text: str
    score: int


def _split_blocks(text: str) -> list[tuple[str, str]]:
    """Split into (nearest_heading, paragraph) pairs, in document order."""
    blocks: list[tuple[str, str]] = []
    heading = ""
    for raw in re.split(r"\n\s*\n", text):
        para = raw.strip()
        if not para:
            continue
        first_line = para.splitlines()[0].strip()
        if _HEADING_RE.match(first_line) and len(first_line) <= 60:
            heading = first_line.lstrip("# ").strip()
            body = "\n".join(para.splitlines()[1:]).strip()
            if body:
                blocks.append((heading, body))
            continue
        blocks.append((heading, para))
    return blocks


def _score_block(heading: str, para: str, index: int) -> int:
    lowered = (heading + "\n" + para).lower()
    score = 0
    score += 3 * sum(1 for kw in _KEYWORDS if kw in lowered)
    if _NUMBER_RE.search(para):
        score += 2
    if any(kw in heading.lower() for kw in _KEYWORDS):
        score += 2
    if index < 10:  # intros usually state the goal and deliverables
        score += 1
    return score


def select_excerpts(
    text: str, source: str, budget: int = PER_FILE_CHAR_BUDGET
) -> list[Excerpt]:
    """Highest-scoring fragments within ``budget`` chars, in document order."""
    scored = [
        (index, Excerpt(source, heading, para, _score_block(heading, para, index)))
        for index, (heading, para) in enumerate(_split_blocks(text))
        if len(para) >= 10
    ]
    picked: list[tuple[int, Excerpt]] = []
    used = 0
    for index, excerpt in sorted(scored, key=lambda pair: (-pair[1].score, pair[0])):
        clipped = excerpt.text[:budget]
        if used + len(clipped) > budget:
            continue
        used += len(clipped)
        picked.append((index, Excerpt(source, excerpt.heading, clipped, excerpt.score)))
        if used >= budget:
            break
    return [excerpt for _, excerpt in sorted(picked, key=lambda pair: pair[0])]
