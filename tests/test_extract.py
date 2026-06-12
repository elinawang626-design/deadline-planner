from pathlib import Path

from planner.extract import extract_text, select_excerpts


def test_plain_text_extraction(tmp_path):
    path = tmp_path / "notes.md"
    path.write_text("# 需求\n\n完成三份交付物，截止 6 月 20 日。", encoding="utf-8")
    status, text = extract_text(path, "notes.md")
    assert status == "ok"
    assert "交付物" in text


def test_gbk_text_falls_back(tmp_path):
    path = tmp_path / "old.txt"
    path.write_bytes("需求：提交报告".encode("gb18030"))
    status, text = extract_text(path, "old.txt")
    assert status == "ok"
    assert "提交报告" in text


def test_unsupported_suffix(tmp_path):
    path = tmp_path / "data.xlsx"
    path.write_bytes(b"binary")
    assert extract_text(path, "data.xlsx") == ("unsupported", None)


def test_corrupt_pdf_marked_failed(tmp_path):
    path = tmp_path / "scan.pdf"
    path.write_bytes(b"not a real pdf")
    status, text = extract_text(path, "scan.pdf")
    assert status == "failed"
    assert text is None


def test_docx_extraction(tmp_path):
    import docx

    document = docx.Document()
    document.add_paragraph("交付要求：实现登录功能")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "验收标准"
    table.rows[0].cells[1].text = "通过 10 个测试"
    path = tmp_path / "spec.docx"
    document.save(str(path))

    status, text = extract_text(path, "spec.docx")
    assert status == "ok"
    assert "交付要求：实现登录功能" in text
    assert "验收标准 | 通过 10 个测试" in text


REQ_DOC = """# 项目说明

这是一段没有信号词的背景介绍，主要讲故事。

## 交付要求

必须提交 3 份文档和 1 个可运行 demo，验收标准是通过全部测试。

## 闲聊

今天天气不错，无关内容。
"""


def test_selection_prefers_requirement_blocks_and_is_stable():
    first = select_excerpts(REQ_DOC, "spec.md")
    second = select_excerpts(REQ_DOC, "spec.md")
    assert first == second  # deterministic
    texts = [e.text for e in first]
    assert any("验收标准" in t for t in texts)
    top = max(first, key=lambda e: e.score)
    assert "必须提交" in top.text
    assert top.heading == "交付要求"


def test_selection_respects_budget():
    long_doc = "\n\n".join(f"需求 {i}：必须完成交付物 {i}，验收截止。" for i in range(100))
    excerpts = select_excerpts(long_doc, "big.md", budget=200)
    assert sum(len(e.text) for e in excerpts) <= 200
